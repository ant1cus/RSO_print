import datetime
import getpass
import os
import pathlib
import re
import shutil
import socket
import time
import traceback
import zipfile

import docx
import fitz
import openpyxl
import pythoncom
import win32api
import win32com
import win32event
from lxml import etree
from win32comext.shell import shell
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns

import win32com.client
import win32print
from PyQt5 import QtPrintSupport
from PyQt5.QtCore import QThread, pyqtSignal
from docx.shared import Pt
from natsort import natsorted
from word2pdf import word2pdf


class PrintDoc(QThread):  # Поток для печати
    progress = pyqtSignal(int)  # Сигнал для progressbar
    status = pyqtSignal(str)  # Сигнал для статус бара
    messageChanged = pyqtSignal(str, str)  # Сигнал, если ошибка

    def __init__(self, incoming_data):  # Список значений для работы потока
        QThread.__init__(self)
        # Присваиваем значения
        self.path_old = incoming_data['path_old_print']
        self.account_num_path = incoming_data['path_account_num']
        self.add_path_account_num = incoming_data['add_path_account_num']
        self.print_flag = incoming_data['print_flag']
        self.name_printer = incoming_data['name_printer']
        self.path_form27 = incoming_data['path_form_27']
        self.print_order = incoming_data['print_order']
        self.service = incoming_data['service']
        self.path_for_def = incoming_data['path_for_default']
        self.logging = incoming_data['logging']
        self.package = incoming_data['package_']
        self.document_list = incoming_data['document_list']
        self.incoming = {info: incoming_data[info] if incoming_data[info] else '' for info in incoming_data}

    def run(self):

        def print_doc(path_old, account_num_path, add_path_account_num, print_flag, name_printer, path_form27,
                      print_order, service, path_for_def, logging, status, progress):

            def rm(folder_path):
                try:
                    while len(os.listdir(folder_path)) != 0:
                        time.sleep(0.5)
                        for file_object in os.listdir(folder_path):
                            flag_ = True
                            while flag_:
                                try:
                                    file_object_path = os.path.join(folder_path, file_object)
                                    try:
                                        if os.path.isfile(file_object_path) or os.path.islink(file_object_path):
                                            os.remove(file_object_path)
                                        else:
                                            try:
                                                shutil.rmtree(file_object_path)
                                            except FileNotFoundError:
                                                pass
                                    except OSError:
                                        os.remove(folder_path)
                                    flag_ = False
                                except BaseException:
                                    pass
                except NotADirectoryError:
                    os.remove(folder_path)
                time.sleep(0.05)
                shutil.rmtree(folder_path)

            # def list_count(name_doc):
            #     pythoncom.CoInitializeEx(0)
            #     name_doc_start = os.path.abspath(os.getcwd() + '\\' + name_doc)
            #     name_doc_file_pdf = name_doc_start + '.pdf'
            #     self.logging.info('Конвертируем в пдф ' + name_doc_start)
            #     docx2pdf.convert(name_doc_start, name_doc_file_pdf)
            #     doc_file_pdf = fitz.open(name_doc_file_pdf)  # Открываем пдф
            #     doc_page = doc_file_pdf.page_count  # Получаем кол-во страниц
            #     doc_file_pdf.close()  # Закрываем
            #     self.logging.info('Удаляем пдф ' + name_doc_start)
            #     os.remove(name_doc_file_pdf)  # Удаляем пдф документ
            #     self.logging.info('Вставляем страницы в word ' + name_doc_start)
            #     temp_docx_ = name_doc_start
            #     temp_zip_ = name_doc_start + ".zip"
            #     temp_folder_ = os.path.join(os.getcwd() + '\\', "template")
            #     os.rename(temp_docx_, temp_zip_)
            #     os.mkdir(os.getcwd() + '\\zip')
            #     with zipfile.ZipFile(temp_zip_) as my_document_:
            #         my_document_.extractall(temp_folder_)
            #     pages_xml_ = os.path.join(temp_folder_, "docProps", "app.xml")
            #     string_ = open(pages_xml_, 'r', encoding='utf-8').read()
            #     string_ = re.sub(r"<Pages>(\w*)</Pages>",
            #                      "<Pages>" + str(doc_page) + "</Pages>", string_)
            #     with open(pages_xml_, "wb") as file_wb_:
            #         file_wb_.write(string_.encode("UTF-8"))
            #     self.logging.info('Получаем word из зип ' + name_doc_start)
            #     os.remove(temp_zip_)
            #     shutil.make_archive(temp_zip_.replace(".zip", ""), 'zip', temp_folder_)
            #     os.rename(temp_zip_, temp_docx_)  # rename zip file to docx
            #     rm(temp_folder_)
            #     rm(os.getcwd() + '\\zip')
            #     return doc_page

            try:  # Ловим ошибку чтобы программа не вылетала молча
                # Проверка на количество листов и учетных номеров
                num_of_sheets = 0
                os.chdir(path_old)  # Меняем рабочую директорию
                percent_val = 0  # Отсылаемое значение в прогресс бар
                docs = [i for i in os.listdir() if i[-4:] == 'docx' and '~' not in i]  # Список файлов
                status.emit('Создаем второй сопроводительный документ')
                logging.info('Второй сопроводительный')
                for el in docs:  # Для второго сопроводительного
                    # if re.findall('сопровод', el.lower()) or re.findall('запрос', el.lower()):
                    if re.findall('запрос', el.lower()):
                        shutil.copy(el, el.rpartition('.')[0] + ' (2 экз.).docx', follow_symlinks=True)
                        doc = docx.Document(os.getcwd() + '\\' + el.rpartition('.')[0] + ' (2 экз.).docx')
                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'TimesNewRoman'
                        font.size = Pt(11)
                        header = doc.sections[0].first_page_header  # Верхний колонтитул первой страницы
                        head = header.paragraphs[0]  # Параграф
                        if re.findall(r'экз.', head.text.lower()):
                            header_text = ''
                            for enum, paragraph in enumerate(doc.sections[0].first_page_header.paragraphs):
                                header_text = doc.sections[0].first_page_header.paragraphs[enum].text
                                header_text = re.sub(r'№1', '№2', header_text)
                                if 'экз.' in paragraph.text.lower():
                                    break
                            doc.save(os.path.abspath(os.getcwd() + '\\' + el.rpartition('.')[0] + ' (2 экз.).docx'))
                            temp_docx = os.path.join(os.getcwd() + '\\' + el.rpartition('.')[0] + ' (2 экз.).docx')
                            temp_zip = os.path.join(os.getcwd() + '\\', el.rpartition('.')[0] + ' (2 экз.).docx.zip')
                            temp_folder = os.path.join(os.getcwd() + '\\', "template")
                            if os.path.exists(temp_zip):
                                shutil.rmtree(temp_zip)
                            if os.path.exists(temp_folder):
                                shutil.rmtree(temp_folder)
                            if os.path.exists(os.getcwd() + '\\zip'):
                                shutil.rmtree(os.getcwd() + '\\zip')
                            os.rename(temp_docx, temp_zip)
                            os.mkdir(os.getcwd() + '\\zip')
                            with zipfile.ZipFile(temp_zip) as my_document:
                                my_document.extractall(temp_folder)
                            name_header = ''
                            size_header = 0
                            for header_file in os.listdir(pathlib.Path(temp_folder, 'word')):
                                if 'header' in header_file:
                                    if os.path.getsize(pathlib.Path(temp_folder, 'word', header_file)) > size_header:
                                        size_header = os.path.getsize(pathlib.Path(temp_folder, 'word', header_file))
                                        name_header = header_file
                            shutil.copy(pathlib.Path(path_for_def, 'documents', 'header1.xml'),
                                        pathlib.Path(temp_folder, 'word', name_header))
                            # shutil.copy(pathlib.Path(path_for_def, 'documents', 'header1.xml'),
                            #             pathlib.Path(temp_folder, 'word', 'header1.xml'))
                            os.remove(temp_zip)
                            shutil.make_archive(temp_zip.replace(".zip", ""), 'zip', temp_folder)
                            os.rename(temp_zip, temp_docx)  # rename zip file to docx
                            while True:
                                try:
                                    shutil.rmtree(temp_folder)
                                    shutil.rmtree(os.getcwd() + '\\zip')
                                    break
                                except OSError as es:
                                    self.logging.error(es)
                                    self.logging.error(traceback.format_exc())
                                    self.logging.info('Ошибка с удалением, пробуем ещё раз')
                            doc = docx.Document(os.getcwd() + '\\' + el.rpartition('.')[0] + ' (2 экз.).docx')
                            header = doc.sections[0].first_page_header  # Верхний колонтитул первой страницы
                            head = header.paragraphs[0]  # Параграф
                            head.text = header_text
                            for header_styles in head.runs:
                                header_styles.font.size = Pt(11)
                                header_styles.font.name = 'Times New Roman'
                            head_format = head.paragraph_format  # Настройки параграфа
                            head_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Выравниваем по правому краю
                            doc.save(os.path.abspath(os.getcwd() + '\\' + el.rpartition('.')[0] + ' (2 экз.).docx'))
                            pythoncom.CoInitializeEx(0)
                            count_file = os.path.abspath(os.getcwd() + '\\' + el.rpartition('.')[0] + ' (2 экз.).docx')
                            name_file_pdf = count_file + '.pdf'
                            self.logging.info('Конвертируем в пдф ' + count_file)
                            word2pdf(str(pathlib.Path(count_file)), str(pathlib.Path(count_file, name_file_pdf)))
                            input_file_pdf = fitz.open(name_file_pdf)  # Открываем пдф
                            count_page = input_file_pdf.page_count  # Получаем кол-во страниц
                            input_file_pdf.close()  # Закрываем
                            self.logging.info('Удаляем пдф ' + count_file)
                            os.remove(name_file_pdf)  # Удаляем пдф документ
                            status.emit('Считаем количество страниц во втором сопроводительном документе')
                            self.logging.info('Вставляем страницы в word ' + count_file)
                            temp_docx = count_file
                            temp_zip = count_file + ".zip"
                            temp_folder = os.path.join(os.getcwd() + '\\', "template")
                            os.rename(temp_docx, temp_zip)
                            os.mkdir(os.getcwd() + '\\zip')
                            with zipfile.ZipFile(temp_zip) as my_document:
                                my_document.extractall(temp_folder)
                            pages_xml = os.path.join(temp_folder, "docProps", "app.xml")
                            string = open(pages_xml, 'r', encoding='utf-8').read()
                            string = re.sub(r"<Pages>(\w*)</Pages>",
                                            "<Pages>" + str(count_page) + "</Pages>", string)
                            with open(pages_xml, "wb") as file_wb:
                                file_wb.write(string.encode("UTF-8"))
                            self.logging.info('Получаем word из зип ' + count_file)
                            os.remove(temp_zip)
                            shutil.make_archive(temp_zip.replace(".zip", ""), 'zip', temp_folder)
                            os.rename(temp_zip, temp_docx)  # rename zip file to docx
                            rm(temp_folder)
                            rm(os.getcwd() + '\\zip')
                docs = [i for i in os.listdir() if i[-4:] == 'docx' and '~' not in i]  # Список файлов
                logging.info('Сортируем')
                docs = natsorted(docs, key=lambda y: y.rpartition(' ')[2][:-5])
                if print_order:
                    logging.info('Есть порядок печати')
                    quantity_docs = {'Заключение': 0, 'Протокол': 0, 'Приложение А': 0, 'Предписание': 0}
                    docs_name = {}
                    for element in docs:
                        for doc_name in ['Заключение', 'Протокол', 'Приложение А', 'Предписание']:
                            if re.findall(doc_name.lower(), element.lower()):
                                quantity_docs[doc_name] += 1
                                doc_number = element.rpartition('.')[0].rpartition(' ')[2]
                                if doc_number in docs_name:
                                    docs_name[doc_number].append(element)
                                else:
                                    docs_name[doc_number] = [element]
                    quantity = flag = 0
                    for element in quantity_docs:
                        if quantity_docs.get(element) != 0:
                            quantity += quantity_docs.get(element)
                            flag += 1
                    if quantity % flag != 0:
                        status.emit('Разное количество документов')
                        return
                    docs_ = []
                    for element in docs_name:
                        for doc_name in ['Заключение', 'Протокол', 'Приложение А', 'Предписание']:
                            for i in docs_name[element]:
                                if re.findall(doc_name.lower(), i.lower()):
                                    docs_.append(i)
                                    break
                    docs_sec = [j for i in ['Форма 3', 'Опись',
                                            'Сопровод'] for j in docs if re.findall(i.lower(), j.lower())]
                    docs_ = docs_ + docs_sec
                else:
                    logging.info('Нет порядка печати')
                    docs_ = [j for i in ['Заключение', 'Протокол', 'Приложение А', 'Предписание', 'Форма 3', 'Опись',
                                         'Сопровод'] for j in docs if re.findall(i.lower(), j.lower())]
                docs_not = [i for i in docs if i not in docs_]
                docs = docs_not + docs_
                logging.info('Отсортированные документы:\n' + '-|-'.join(docs))
                per = 0
                for file in docs:
                    if file.partition(' ')[0].lower() not in ['протокол', 'приложение а'] and service:
                        per += 1
                try:
                    percent = 100 / per  # Процент от общего
                except ZeroDivisionError:
                    percent = 0
                    logging.warning('Деление на 0, ни одного документа для печати')
                status.emit('Считаем количество листов в документах...')
                logging.info('Считаем количество листов в документах')

                def list_doc(dp):
                    # Открываем word документ как зип архив для доступа к xml свойствам
                    with zipfile.ZipFile(os.getcwd() + '\\' + dp) as my_doc:
                        xml_content = my_doc.read('docProps/app.xml')  # Общие свойства
                        pages_ = re.findall(r'<Pages>(\w*)</Pages>', xml_content.decode())  # Ищем кол-во страниц
                        if int(pages_[0]) > 1:
                            ns = int(pages_[0]) - 1
                        else:
                            ns = int(pages_[0])
                            # ns = list_count(dp)  # Для проверки, вдруг изменилось количество страниц
                    return ns
                num_of_sheets_logs = {}
                for doc_path in docs:
                    if doc_path.endswith('.docx'):
                        if re.findall('заключение', doc_path.lower()) and self.document_list['заключение'] is False:
                            continue
                        elif re.findall('протокол', doc_path.lower()) and self.document_list['протокол'] is False:
                            continue
                        elif re.findall('предписание', doc_path.lower()) and self.document_list['предписание'] is False:
                            continue
                        elif re.findall('приложение а', doc_path.lower()):
                            continue
                        else:
                            num_of_sheets_logs[doc_path] = list_doc(doc_path)
                            num_of_sheets += num_of_sheets_logs[doc_path]
                        # if service:
                        #     if re.findall('заключение', doc_path.lower()) and self.document_list['заключение'] is False:
                        #         continue
                        #     elif re.findall('протокол', doc_path.lower()) and self.document_list['протокол'] is False:
                        #         continue
                        #     elif re.findall('предписание', doc_path.lower())\
                        #             and self.document_list['предписание'] is False:
                        #         continue
                        #     else:
                        #         num_of_sheets_logs[doc_path] = list_doc(doc_path)
                        #         num_of_sheets += num_of_sheets_logs[doc_path]
                        # else:
                        #     if re.findall('заключение', doc_path.lower()) and self.document_list['заключение'] is False:
                        #         continue
                        #     elif re.findall('протокол', doc_path.lower()) and self.document_list['протокол'] is False:
                        #         continue
                        #     elif re.findall('предписание', doc_path.lower())\
                        #             and self.document_list['предписание'] is False:
                        #         continue
                        #     elif re.findall('приложение', doc_path.lower()):
                        #         continue
                        #     else:
                        #         num_of_sheets_logs[doc_path] = list_doc(doc_path)
                        #         num_of_sheets += num_of_sheets_logs[doc_path]
                logging.info('Листы в документах:')
                logging.info(num_of_sheets_logs)

                def numbers_list(path_s, acc_num, num_in, num_del=-1):  # Для подсчета кол-ва номеров в файле с номерами
                    w_b = openpyxl.load_workbook(path_s)  # Открываем книгу
                    w_s = w_b.active  # Делаем активный лист
                    flag_ = 0  # Флаг для выхода
                    for j_ in range(1, w_s.max_column + 1):  # По столбцам
                        for i_ in range(1, w_s.max_row + 1):  # По строкам
                            if w_s.cell(i_, j_).value:  # Если есть значение
                                if num_del > -1:  # Если на удаление
                                    num_del += 1  # Подсчет
                                    w_s.cell(i_, j_).value = ''  # Удаляем значение
                                    if num_del == num_in:  # Если подсчет равен кол-ву требуемых листов
                                        flag_ = 1  # Для дальнейшего выхода
                                        break  # Выход
                                else:  # Если на подсчет кол-ва номеров
                                    num_in += 1  # Подсчет
                                    acc_num.append(w_s.cell(i_, j_).value)  # Присваиваем значение
                                    if num_of_sheets == num_in:  # Если значений столько, сколько нужно
                                        flag_ = 1  # Для дальнейшего выхода
                                        break  # Выход
                        if flag_:  # Если нужен выход
                            break  # Выход
                    w_b.save(path_s)  # Сохраняем
                    w_b.close()  # Закрываем
                    if num_del > -1:  # Если для удаления
                        return num_del  # Возвращаем значения для удаления
                    else:  # Если для подсчета
                        return num_in  # Возвращаем значение для подсчета

                logging.info('Номера для печати')
                acc_num_for_print = []  # Номера для печати
                num_in_file = 0  # Подсчет кол-ва
                num_in_file = numbers_list(account_num_path, acc_num_for_print, num_in_file)  # В функцию
                if add_path_account_num:  # Если есть доп. файл
                    num_in_file = numbers_list(add_path_account_num, acc_num_for_print, num_in_file)  # В функцию
                num_for_del = 0  # Подсчет для удаления
                if num_of_sheets <= num_in_file:  # Если номеров хватает
                    num_for_del = numbers_list(account_num_path, acc_num_for_print, num_in_file, num_for_del)  # В ф-ию
                    if add_path_account_num:  # Если есть доп. файл, то в ф-ию
                        numbers_list(add_path_account_num, acc_num_for_print, num_in_file, num_for_del)
                    status.emit('Листы посчитаны')
                else:  # Если номеров не хватает
                    for document in os.listdir(path_old):
                        if re.findall(r'сопровод', document.lower()) or re.findall(r'запрос', document.lower()):
                            if re.findall(r'экз', document.lower()):
                                os.remove(path_old + '\\' + document)
                    status.emit('Не хватает номеров учетных листов, загрузите дополнительный файл!')
                    return 'Не хватает номеров учетных листов, загрузите дополнительный файл!'
                # print('Number', len(acc_num_for_print))
                # print(acc_num_for_print)

                def del_col(path_save):  # Ф-я для удаления пустых колонок в файле с учетными номерами листов
                    w_b = openpyxl.load_workbook(path_save)  # Открываем
                    w_s = w_b.active  # Активный лист
                    flag_ = False  # Для выхода
                    for j_ in range(1, w_s.max_column + 1):  # Колонки
                        for i_ in range(1, w_s.max_row + 1):  # Строки
                            if w_s.cell(i_, j_).value:  # Если есть значение
                                flag_ = True  # Для дальнейшего выхода
                                break  # Выход
                        if flag_:  # Если есть метка
                            break  # Выход
                        w_s.delete_cols(1, j_)  # Если прошли всю колонку и не вышли - удаляем колонку
                    w_b.save(path_save)  # Сохраняем книгу
                    w_b.close()  # Закрываем

                logging.info('Удаляем используемые номера')
                del_col(account_num_path)  # Удаляем колонку в файле
                if add_path_account_num:  # Если есть доп. файл номеров
                    del_col(add_path_account_num)  # Удаляем колонку

                num_for_print = 0
                win32print.SetDefaultPrinter(name_printer)
                user_name = getpass.getuser()
                printer = QtPrintSupport.QPrinterInfo.defaultPrinterName()
                computer_name = socket.gethostname()
                logging.info('Номера для печати')
                logging.info(len(acc_num_for_print))
                logging.info(acc_num_for_print)
                logging.info('Кол-во листов')
                logging.info(num_of_sheets)
                for el in docs:  # Для файлов в папке
                    flag_for_exit = True
                    name_pdf = el.rpartition('.')[0] + '.pdf'  # Путь для пдф файла
                    while flag_for_exit:
                        try:
                            pythoncom.CoInitializeEx(0)
                            logging.info('Форматируем документ ' + str(el))
                            status.emit('Форматируем документ ' + str(el))
                            if re.findall(r'приложение а', el.lower()):
                                if service is True:
                                    logging.info('Не печатаем приложение ' + str(el) + ' (service true)')
                                    flag_for_exit = False
                                    continue
                                status.emit('Печатаем документ ' + str(el))
                                logging.info('Запускаем в печать ' + str(el))
                                printing_date = [computer_name, user_name, path_old + '\\' + str(el),
                                                 str(datetime.date.today()), printer]
                                win32api.ShellExecute(0, "print", path_old + '\\' + el,
                                                      name_printer, ".", 0)
                                jobs = 0  # Проверка для того, что бы не перескакивать на следующий документ
                                printer_defaults = {"DesiredAccess": win32print.PRINTER_ACCESS_USE}  # Дефолтный принтер
                                handle = win32print.OpenPrinter(name_printer, printer_defaults)  # Открываем
                                logging.info('Ждем очередь ' + str(el))
                                while jobs < 3:
                                    print_jobs = win32print.EnumJobs(handle, 0, -1, 1)  # Очередь печати
                                    if not print_jobs and jobs == 0:  # Пока не запустилось в печать
                                        pass
                                    elif not print_jobs and jobs == 2:  # Если запустилось и очистилась
                                        jobs = 3
                                        logging.info('Очередь очистилась')
                                    elif print_jobs:  # Если в очереди что-то есть
                                        jobs = 2
                            else:
                                if re.findall('заключение', el.lower()) and self.document_list['заключение'] is False:
                                    flag_for_exit = False
                                    continue
                                elif re.findall('протокол', el.lower()) and self.document_list['протокол'] is False:
                                    flag_for_exit = False
                                    continue
                                elif re.findall('предписание', el.lower()) and\
                                        self.document_list['предписание'] is False:
                                    flag_for_exit = False
                                    continue
                                num_start = acc_num_for_print[num_for_print]
                                num_second_page = '' if num_of_sheets_logs[el] == 1 \
                                    else acc_num_for_print[num_for_print + 1]
                                num_for_print += num_of_sheets_logs[el]
                                num_stop = acc_num_for_print[num_for_print - 1]
                                logging.info('Вставляем номера листов ' + str(el))

                                def create_element(attrib_name):
                                    return OxmlElement(attrib_name)

                                def create_attribute(attrib, attrib_name, attrib_value):
                                    attrib.set(ns.qn(attrib_name), attrib_value)

                                def add_page_number(paragraph, value_num, number_page=''):
                                    # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                    # if orientation is False:
                                    page_run = paragraph.add_run()
                                    t1 = create_element('w:t')
                                    create_attribute(t1, 'xml:space', 'preserve')
                                    t1.text = '\t\t' + value_num
                                    page_run._r.append(t1)

                                    page_num_run = paragraph.add_run()

                                    fld_char1 = create_element('w:fldChar')
                                    create_attribute(fld_char1, 'w:fldCharType', 'begin')

                                    instr_text_or1 = create_element('w:instrText')
                                    create_attribute(instr_text_or1, 'xml:space', 'preserve')
                                    instr_text_or1.text = "="

                                    fld_char2 = create_element('w:fldChar')
                                    create_attribute(fld_char2, 'w:fldCharType', 'begin')

                                    instrText = create_element('w:instrText')
                                    create_attribute(instrText, 'xml:space', 'preserve')
                                    instrText.text = "PAGE"

                                    fld_char3 = create_element('w:fldChar')
                                    create_attribute(fld_char3, 'w:fldCharType', 'end')

                                    instr_text_or2 = create_element('w:instrText')
                                    create_attribute(instr_text_or2, 'xml:space', 'preserve')
                                    instr_text_or2.text = " - 2 +" + number_page

                                    fld_char4 = create_element('w:fldChar')
                                    create_attribute(fld_char4, 'w:fldCharType', 'end')

                                    page_num_run._r.append(fld_char1)
                                    page_num_run._r.append(instr_text_or1)
                                    page_num_run._r.append(fld_char2)
                                    page_num_run._r.append(instrText)
                                    page_num_run._r.append(fld_char3)
                                    page_num_run._r.append(instr_text_or2)
                                    page_num_run._r.append(fld_char4)

                                doc = docx.Document(pathlib.Path(path_old, el))  # Открываем
                                footer_1 = doc.sections[0].first_page_footer  # Нижний колонтитул первой страницы
                                foot_1 = footer_1.paragraphs[0]  # Параграф
                                foot_1.text = footer_1.paragraphs[0].text + '\t\t' + num_start  # Текст
                                foot_format = foot_1.paragraph_format  # Настройки параграфа
                                foot_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выравнивание по левому краю
                                if num_second_page:
                                    footer_2 = doc.sections[1].footer.paragraphs[0]  # Нижний колонтитул страницы
                                    mask_page = num_second_page.rpartition('/')[0] + '/'
                                    start_number = num_second_page.rpartition('/')[2]
                                    add_page_number(footer_2, mask_page, start_number)
                                doc.save(pathlib.Path(path_old, el))  # Сохраняем
                                # print('doc', el)
                                # print('start', num_start, 'stop', num_stop)
                                # logging.info('Преобразуем в pdf ' + str(el))
                                try:
                                    word2pdf(str(pathlib.Path(path_old, el)), str(pathlib.Path(path_old, name_pdf)))
                                except BaseException:
                                    word = win32com.client.Dispatch("Word.Application")
                                    word.Quit()
                                doc_old = docx.Document(path_old + '\\' + el)  # Открываем
                                last = doc_old.sections[len(doc_old.sections) - 1].first_page_footer  # Колонтитул
                                logging.info('Вставляем номера в 27 форму')
                                number = last.paragraphs[0].text.partition('\n')[0].rpartition(' ')[2]
                                if path_form27:
                                    # Если файл уже был преобразован и название отличается от «Форма 27»
                                    path_form27_dir = pathlib.Path(path_form27).parent
                                    path_form27_file = [i for i in os.listdir(path_form27_dir)
                                                        if 'форма 27' in i.lower()]
                                    wb = openpyxl.open(path_form27_file[0])
                                    ws = wb.active
                                    for row in range(1, ws.max_row):
                                        if ws.cell(row, 1).value == number:
                                            if re.findall('сопровод', el.lower()):
                                                if re.findall(' (2 экз.)', el.lower()):
                                                    ws.cell(row + 3, 11).value = num_start
                                                    if num_start != num_stop:
                                                        ws.cell(row + 4, 11).value = num_stop
                                                    break
                                                else:
                                                    ws.cell(row, 11).value = num_start
                                                    if num_start != num_stop:
                                                        ws.cell(row + 1, 11).value = num_stop
                                                    break
                                            else:
                                                ws.cell(row, 11).value = num_start
                                                if num_start != num_stop:
                                                    ws.cell(row + 1, 11).value = num_stop
                                                break
                                    if re.findall('сопровод', el.lower()):
                                        for row in range(2, ws.max_row):
                                            if ws.cell(row, 1).value:
                                                if ws.cell(row, 1).value == number:
                                                    break
                                                else:
                                                    ws.cell(row, 17).value = 'Уч. ном. ' + number
                                    wb.save(filename=path_form27.rpartition('\\')[0] + '\\' + path_form27_file[0])
                                    wb.close()
                                if print_flag == 'Односторонняя':  # Если печать односторонняя - печатаем
                                    status.emit('Печатаем документ ' + str(el))
                                    logging.info('Печатаем документ ' + str(el))
                                    printing_date = [computer_name, user_name, path_old + '\\' + str(el),
                                                     str(datetime.date.today()), printer]
                                    win32api.ShellExecute(0, "print", path_old + '\\' + name_pdf,
                                                          name_printer, ".", 0)
                                    jobs = 0  # Проверка для того, что бы не перескакивать на следующий документ
                                    # Дефолтный принтер
                                    printer_defaults = {"DesiredAccess": win32print.PRINTER_ACCESS_USE}
                                    handle = win32print.OpenPrinter(name_printer, printer_defaults)  # Открываем
                                    logging.info('Ждем очередь ' + str(el))
                                    while jobs < 3:
                                        print_jobs = win32print.EnumJobs(handle, 0, -1, 1)  # Очередь печати
                                        if not print_jobs and jobs == 0:  # Пока не запустилось в печать
                                            pass
                                        elif not print_jobs and jobs == 2:  # Если запустилось и очистилась
                                            jobs = 3
                                            logging.info('Очередь очистилась ' + str(el))
                                        elif print_jobs:  # Если в очереди что-то есть
                                            jobs = 2
                                else:
                                    if print_flag == 'Двухсторонняя последняя':  # Если последняя страница дуплекс
                                        status.emit('Преобразуем документ ' + str(el))
                                        printing_date = [computer_name, user_name, path_old + '\\' + str(el),
                                                         str(datetime.date.today()), printer]
                                        # Дефолтный принтер
                                        printer_defaults = {"DesiredAccess": win32print.PRINTER_ACCESS_USE}
                                        logging.info('Преобразуем документ ' + str(el))
                                        handle = win32print.OpenPrinter(name_printer, printer_defaults)  # Открываем
                                        input_file = fitz.open(path_old + '/' + name_pdf)  # Открываем пдф
                                        pages = input_file.page_count  # Получаем кол-во страниц
                                        if pages == 2:
                                            level = 2
                                            # Получаем атрибуты принтера
                                            attributes = win32print.GetPrinter(handle, level)
                                            attributes['pDevMode'].Duplex = 2  # flip up  Для двухсторонней печати
                                            try:
                                                # Устанавливаем настройки
                                                win32print.SetPrinter(handle, level, attributes, 0)
                                            except:  # Пропускаем ошибку
                                                pass
                                            # Печатаем
                                            logging.info('Печатаем ' + name_pdf)
                                            win32api.ShellExecute(0, 'print', path_old + '\\' + name_pdf,
                                                                  name_printer, '.', 0)
                                            status.emit('Печатаем документ ' + name_pdf)
                                            jobs = 0  # Проверка для того, что бы не перескакивать на следующий документ
                                            logging.info('Ждем очередь ' + name_pdf)
                                            while jobs < 3:
                                                print_jobs = win32print.EnumJobs(handle, 0, -1, 1)  # Очередь печати
                                                if not print_jobs and jobs == 0:  # Пока не запустилось в печать
                                                    pass
                                                elif not print_jobs and jobs == 2:  # Если запустилось и очистилась
                                                    jobs = 3
                                                    logging.info('Очередь очистилась')
                                                elif print_jobs:  # Если в очереди что-то есть
                                                    jobs = 2
                                            input_file.close()
                                        else:
                                            logging.info('Преобразуем документ ' + name_pdf)
                                            output_1_side = path_old + '/' + '1_' + name_pdf  # Путь для сохранения пдф
                                            # Страницы для односторонней печати
                                            selected_page = [page for page in range(0, pages - 2)]
                                            input_file.select(selected_page)  # Выбираем страницы
                                            input_file.save(output_1_side)  # Сохраняем файл
                                            # Печатаем
                                            win32api.ShellExecute(0, 'print', output_1_side, name_printer, '.', 0)
                                            input_file = fitz.open(path_old + '\\' + name_pdf)  # Открываем еще раз
                                            output_2_side = path_old + '\\' + '2_' + name_pdf  # Путь для сохранения
                                            selected_page = [pages - 2, pages - 1]  # Страницы для двухсторонней печати
                                            input_file.select(selected_page)  # Выбираем страницы
                                            input_file.save(output_2_side)  # Сохраняем
                                            input_file.close()
                                            jobs = 0  # Проверка для того, что бы не перескакивать на следующий документ
                                            status.emit('Печатаем документ ' + str(el))
                                            logging.info('Ждем очередь')
                                            while jobs < 3:
                                                print_jobs = win32print.EnumJobs(handle, 0, -1, 1)  # Очередь печати
                                                if not print_jobs and jobs == 0:  # Пока не запустилось в печать
                                                    pass
                                                elif not print_jobs and jobs == 2:  # Если запустилось и очистилась
                                                    jobs = 3
                                                    logging.info('Очередь очистилась')
                                                    os.remove(output_1_side)  # Удаляем файл
                                                elif print_jobs:  # Если в очереди что-то есть
                                                    jobs = 2
                                            level = 2
                                            # Получаем атрибуты принтера
                                            attributes = win32print.GetPrinter(handle, level)
                                            attributes['pDevMode'].Duplex = 2  # flip up  Для двухсторонней печати
                                            logging.info('Меняем настройки принтера ')
                                            try:
                                                # Устанавливаем настройки
                                                win32print.SetPrinter(handle, level, attributes, 0)
                                            except:  # Пропускаем ошибку
                                                pass
                                            # Печатаем
                                            logging.info('Печатаем')
                                            win32api.ShellExecute(0, 'print', output_2_side, name_printer, '.', 0)
                                            jobs = 0  # Проверка для того, что бы не перескакивать на следующий документ
                                            status.emit('Печатаем последнюю страницу документа ' + str(el))
                                            logging.info('Ждем очередь')
                                            while jobs < 3:
                                                print_jobs = win32print.EnumJobs(handle, 0, -1, 1)  # Очередь печати
                                                if not print_jobs and jobs == 0:  # Пока не запустилось в печать
                                                    pass
                                                elif not print_jobs and jobs == 2:  # Если запустилось и очистилась
                                                    jobs = 3
                                                    logging.info('Очередь очистилась')
                                                    os.remove(output_2_side)  # Удаляем файл
                                                elif print_jobs:  # Если в очереди что-то есть
                                                    jobs = 2
                                    else:
                                        status.emit('Печатаем документ ' + str(el))
                                        # Дефолтный принтер
                                        logging.info('Меняем настройки принтера')
                                        printer_defaults = {"DesiredAccess": win32print.PRINTER_ACCESS_USE}
                                        handle = win32print.OpenPrinter(name_printer, printer_defaults)  # Открываем
                                        level = 2
                                        attributes = win32print.GetPrinter(handle, level)  # Получаем атрибуты принтера
                                        attributes['pDevMode'].Duplex = 2  # flip up  Для двухсторонней печати
                                        try:
                                            win32print.SetPrinter(handle, level, attributes,
                                                                  0)  # Устанавливаем настройки
                                        except:  # Пропускаем ошибку
                                            pass
                                        # Печатаем
                                        printing_date = [computer_name, user_name, path_old + '\\' + str(el),
                                                         str(datetime.date.today()), printer]
                                        logging.info('Печатаем ' + str(el))
                                        win32api.ShellExecute(0, 'print', path_old + '\\' + name_pdf,
                                                              name_printer, '.', 0)
                                        jobs = 0  # Проверка для того, что бы не перескакивать на следующий документ
                                        logging.info('Ждем очередь ' + str(el))
                                        while jobs < 3:
                                            print_jobs = win32print.EnumJobs(handle, 0, -1, 1)  # Очередь печати
                                            if not print_jobs and jobs == 0:  # Пока не запустилось в печать
                                                pass
                                            elif not print_jobs and jobs == 2:  # Если запустилось и очистилась
                                                jobs = 3
                                                logging.info('Очередь очистилась')
                                            elif print_jobs:  # Если в очереди что-то есть
                                                jobs = 2
                                    status.emit('Документ ' + str(el) + ' готов')
                                    logging.info('Меняем настройки принтера')
                                    attributes['pDevMode'].Duplex = 1  # Настройки по умолчанию (односторонняя печать)
                                    try:
                                        win32print.SetPrinter(handle, level, attributes, 0)  # Выставляем настройки
                                    except:
                                        pass
                                    win32print.ClosePrinter(handle)  # Закрываем принтер
                            logging.info('Записываем данные с печати')
                            date_for_saving = datetime.date.today()
                            if os.path.exists(path_for_def / (str(date_for_saving) + '.txt')):
                                with open(path_for_def / (str(date_for_saving) + '.txt'), 'a') as f:
                                    f.write(';'.join(printing_date) + '\n')
                            else:
                                with open(path_for_def / (str(date_for_saving) + '.txt'), 'w') as f:
                                    f.write(';'.join(printing_date) + '\n')
                            flag_for_exit = False
                        except BaseException:
                            logging.info(traceback.format_exc())
                    if os.path.exists(pathlib.Path(path_old, name_pdf)):
                        logging.info('Удаляем пдф ' + name_pdf)
                        os.remove(path_old + '\\' + name_pdf)
                    # if re.findall(r'сопровод', el.lower()) or re.findall(r'запрос', el.lower()):
                    if re.findall(r'запрос', el.lower()):
                        if re.findall(r'экз', el.lower()):
                            os.remove(path_old + '\\' + el)
                    if el.partition(' ')[0].lower() in ['протокол', 'приложение'] and service:
                        pass
                    else:
                        pass
                        percent_val += percent  # Увеличиваем прогресс
                    progress.emit(int(percent_val))  # Посылаем значение в прогресс бар
            except Exception as e:  # Если ошибка
                for document in os.listdir(path_old):
                    if re.findall(r'сопровод', document.lower()) or re.findall(r'запрос', document.lower()):
                        if re.findall(r'экз', document.lower()):
                            os.remove(path_old + '\\' + document)
                self.status.emit('Ошибка')  # Сообщение в статус бар
                self.logging.error("Ошибка:\n " + str(e) + '\n' + traceback.format_exc())

        self.logging.info("\n***********************************************************************************\n")
        self.logging.info("Новый запуск")
        time_start = datetime.datetime.now()
        self.progress.emit(0)  # Обнуление прогресс бара
        self.status.emit('Начинаем печать документов')
        if self.package:
            for folder in os.listdir(self.path_old):
                path_ = self.path_old + '\\' + folder
                path_form27_ = path_ + '\\' + 'Форма 27.xlsx' if self.path_form27 else False
                self.incoming['path_old_print'], self.incoming['path_form_27'] = path_, path_form27_
                self.logging.info('Входные параметры:')
                self.logging.info(self.incoming)
                ex = print_doc(path_, self.account_num_path, self.add_path_account_num, self.print_flag,
                               self.name_printer, path_form27_, self.print_order, self.service, self.path_for_def,
                               self.logging, self.status, self.progress)
                if ex:
                    # print(ex)
                    self.messageChanged.emit("ВНИМАНИЕ!", ex)
                    return
        else:
            self.logging.info('Входные параметры:')
            self.logging.info(self.incoming)
            ex = print_doc(self.path_old, self.account_num_path, self.add_path_account_num, self.print_flag,
                           self.name_printer, self.path_form27, self.print_order, self.service, self.path_for_def,
                           self.logging, self.status, self.progress)
            if ex:
                # print(ex)
                self.messageChanged.emit("ВНИМАНИЕ!", ex)
                return
        self.progress.emit(100)  # Завершаем прогресс бар
        self.logging.info("Конец программы, время работы: " + str(datetime.datetime.now() - time_start))
        self.logging.info("\n***********************************************************************************\n")
        self.status.emit('Печать документов завершена')  # Сообщение в статус бар
