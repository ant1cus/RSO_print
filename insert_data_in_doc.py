from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Pt
import pandas as pd


def insert_header(document: pd.Series):
	if re.findall(r'приложение а', name_el.lower()):
		number_protocol = name_el.rpartition(' ')[2].rpartition('.')[0]
		for appendix_num in for_27:
			if re.findall('протокол', appendix_num[4].lower()) \
					and re.findall(number_protocol, appendix_num[4]):
				for p in doc.paragraphs:
					if re.findall(r'date', p.text):
						text = re.sub(r'date', 'к протоколу уч. № ' + appendix_num[0] + ' от date',
									  p.text)
						p.text = text
						for run in p.runs:
							run.font.size = Pt(pt_num)
							run.font.name = 'Times New Roman'
						break
				break
	if conclusion_num_text:
		for val_p, p in enumerate(doc.paragraphs):
			if re.findall(r'\[ЗАКЛНОМ]', p.text):
				text = re.sub(r'\[ЗАКЛНОМ]', conclusion_num_text, p.text)
				p.text = text
				for run in p.runs:
					run.font.size = Pt(12)
					run.font.name = 'Times New Roman'
				break
	for p in doc.paragraphs:
		if re.findall(r'\[АКТНОМ\]', p.text):
			text = re.sub(r'\[АКТНОМ\]', 'от date № ' + act_number, p.text)
			p.text = text
			for run in p.runs:
				run.font.size = Pt(12)
				run.font.name = 'Times New Roman'
			break
	for val_p, p in enumerate(doc.paragraphs):
		if re.findall(r'\[ЗАКЛНОМ]', p.text):
			text = re.sub(r'\[ЗАКЛНОМ]', conclusion_num_text, p.text)
			p.text = text
			doc.paragraphs[val_p].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
			for run in p.runs:
				run.font.size = Pt(pt_num)
				run.font.name = 'Times New Roman'
			break
	if conclusion_num_text or protocol_num_text:
		for val_p, p in enumerate(doc.paragraphs):
			if conclusion_num_text:
				if re.findall(r'\[ЗАКЛНОМ]', p.text):
					text = re.sub(r'\[ЗАКЛНОМ]', conclusion_num_text, p.text)
					p.text = text
					doc.paragraphs[val_p].paragraph_format.alignment = \
						WD_PARAGRAPH_ALIGNMENT.JUSTIFY
					for run in p.runs:
						run.font.bold = False
						run.font.size = Pt(pt_num)
						run.font.name = 'Times New Roman'
					break_flag[0] = True
			else:
				break_flag[0] = True
			if protocol_num_text:
				if re.findall(r'\[ПРОТНОМ]', p.text):
					text = re.sub(r'\[ПРОТНОМ]', protocol_num_text, p.text)
					p.text = text
					doc.paragraphs[val_p].paragraph_format.alignment = \
						WD_PARAGRAPH_ALIGNMENT.JUSTIFY
					for run in p.runs:
						run.font.size = Pt(pt_num)
						run.font.name = 'Times New Roman'
					break_flag[1] = True
			else:
				break_flag[1] = True
			if all(break_flag):
				break

	header_1 = doc_.sections[0].first_page_header  # Верхний колонтитул первой страницы
	head_1 = header_1.paragraphs[0]  # Параграф
	head_1.insert_paragraph_before(text_first_header_)  # Вставляем перед колонтитулом
	head_1 = header_1.paragraphs[0]  # Выбираем новый первый параграф
	for header_styles in head_1.runs:
		header_styles.font.size = Pt(pt_count)
		header_styles.font.name = 'Times New Roman'
	head_1_format = head_1.paragraph_format  # Настройки параграфа
	head_1_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Выравниваем по правому краю
	footer_ = doc_.sections[0].first_page_footer  # Нижний колонтитул первой страницы
	foot_ = footer_.paragraphs[0]  # Параграф
	foot_.text = text_for_foot_  # Текст
	foot_format_ = foot_.paragraph_format  # Настройки параграфа
	foot_format_.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выравнивание по левому краю
	doc_.sections[0].footer.paragraphs[0].text = text_for_foot_  # Номера для страниц
	# Выравниваем по левому краю
	doc_.sections[0].footer.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
	doc_.add_section()  # Добавляем последнюю страницу
	last_ = doc_.sections[len(doc_.sections) - 1].first_page_header  # Колонтитул для последней страницы
	last_.is_linked_to_previous = False  # Отвязываем от предыдущей секции чтобы не повторялись
	foot_ = doc_.sections[len(doc_.sections) - 1].first_page_footer  # Нижний колонтитул
	foot_.is_linked_to_previous = False  # Отвязываем
	# Текст для фонарика
	foot_.paragraphs[0].text = "Уч. № " + text_for_foot_ + \
							   "\nОтп. 1 экз. в адрес\n" + hdd_number_ + \
							   "\nИсп. " + exec_ + "\nПеч. " + print_people_ + "\n" + \
							   date_ + "\nб/ч"
	for footer_style in foot_.paragraphs[0].runs:
		footer_style.font.size = Pt(pt_count)
		footer_style.font.name = 'Times New Roman'
	if fso_:
		if 'заключение' in name_file_.lower() or 'акт' in name_file_.lower():
			path_new = path_new + '\\' + 'Материалы по специальной проверке технических средств'
		else:
			path_new = path_new + '\\' + 'Материалы по специальным исследованиям технических средств'
		try:
			os.mkdir(path_new)
		except FileExistsError:
			pass
	# logging.info("Вставляем номера страниц")
	# header_2 = doc_.sections[1].header.paragraphs[0]  # Колонтитул страницы для номера
	# add_page_number(header_2)
	doc_.save(os.path.abspath(path_new + '\\' + name_file_))  # Сохраняем